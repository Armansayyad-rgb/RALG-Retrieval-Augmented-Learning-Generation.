"""Document upload pipeline for the Gradio web UI.

Parses uploaded PDFs, DOCX, and TXT files, splits them into chunks using
the same logic as the static knowledge base, and merges the new chunks
into the running pipeline's retrieval index.

Design notes:

- We avoid hard dependencies on PyPDF2/python-docx by lazy-importing them
  only when the corresponding file type is uploaded. This keeps the
  basic chat UI working even if those packages are not installed.
- Chunk sizes mirror ``retriever_v2.load_chunks`` so retrieval scoring
  behaves identically on uploaded and built-in content.
- New uploads extend the lexical postings index in O(new chunks). Deletion
  still rebuilds the index because removing arbitrary chunk positions
  requires compacting postings safely.

Provenance limitations (this checkpoint):

- page_number is always None. The current PDF/DOCX parsers do not
  preserve page boundaries during text extraction.  A future checkpoint
  may enhance page-level provenance.
- revision is always None. Real document versioning will be added later.
"""

from __future__ import annotations

import copy
import logging
import json
import os
import re
import tempfile
import unicodedata
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable

import uuid
from threading import RLock


from retriever_v2 import (
    RuntimeChunk,
    build_index as build_index_v2,
    extend_index as extend_index_v2,
)
from config import RUNTIME_UPLOAD_DIR, UPLOAD_POLICY


_LOGGER = logging.getLogger(__name__)

SUPPORTED_EXTS = UPLOAD_POLICY.allowed_extensions
MAX_TXT_SIZE = UPLOAD_POLICY.per_file_bytes[".txt"]
MAX_PDF_SIZE = UPLOAD_POLICY.per_file_bytes[".pdf"]
MAX_DOCX_SIZE = UPLOAD_POLICY.per_file_bytes[".docx"]
MAX_EXTRACTED_TEXT_LEN = UPLOAD_POLICY.max_extracted_text_chars
MAX_UPLOADED_CHUNKS = UPLOAD_POLICY.max_chunks_per_document

# Maximum length for a sanitized display filename
_MAX_DISPLAY_NAME_LEN = 200
_REGISTRY_NAME = "metadata.json"
_LIFECYCLE_LOCK = RLock()


def _persistence_dir(pipeline: dict) -> Path:
    return Path(pipeline.get("runtime_upload_dir", RUNTIME_UPLOAD_DIR))


def _atomic_write(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = None
    try:
        with tempfile.NamedTemporaryFile(
            "w", encoding="utf-8", dir=path.parent, delete=False
        ) as handle:
            temporary = Path(handle.name)
            handle.write(content)
        temporary.replace(path)
    except OSError:
        if temporary is not None and temporary.exists():
            try:
                temporary.unlink()
            except OSError:
                pass
        raise


def _registry_path(pipeline: dict) -> Path:
    return _persistence_dir(pipeline) / _REGISTRY_NAME


def _persist_registry(pipeline: dict, entries: list[dict]) -> None:
    _atomic_write(_registry_path(pipeline), json.dumps(entries, indent=2))


def _persist_document(pipeline: dict, doc: "UploadedDocument") -> None:
    root = _persistence_dir(pipeline)
    documents_dir = root / "documents"
    documents_dir.mkdir(parents=True, exist_ok=True)
    content_path = documents_dir / f"{doc.doc_id}.txt"
    pre_existed = content_path.exists()
    old_content = None
    if pre_existed:
        try:
            old_content = content_path.read_text(encoding="utf-8")
        except OSError:
            old_content = None
    try:
        _atomic_write(content_path, doc.text)
        entries = _load_registry(pipeline)
        entries = [entry for entry in entries if entry.get("document_id") != doc.doc_id]
        entries.append({
            "document_id": doc.doc_id,
            "document_name": doc.safe_display_name,
            "extension": doc.ext,
            "upload_timestamp": doc.upload_timestamp,
            "source_type": doc.source_type,
            "revision": doc.revision,
            "chunk_count": doc.chunk_count,
            "content_file": f"documents/{doc.doc_id}.txt",
        })
        _persist_registry(pipeline, entries)
    except OSError:
        if pre_existed and old_content is not None:
            try:
                _atomic_write(content_path, old_content)
            except OSError:
                pass
        elif not pre_existed and content_path.exists():
            try:
                content_path.unlink()
            except OSError:
                pass
        raise


def _load_registry(pipeline: dict) -> list[dict]:
    path = _registry_path(pipeline)
    if not path.exists():
        return []
    try:
        value = json.loads(path.read_text(encoding="utf-8"))
        if not isinstance(value, list):
            raise ValueError("registry is not a list")
        return [entry for entry in value if isinstance(entry, dict)]
    except Exception:
        _LOGGER.exception("Failed to load runtime document registry")
        return []


def restore_persisted_documents(pipeline: dict) -> list[UploadedDocument]:
    """Load valid persisted documents; skip bad entries without stopping startup."""
    root = _persistence_dir(pipeline)
    documents_dir = root / "documents"
    restored: list[UploadedDocument] = []
    seen: set[str] = set()
    skipped = 0
    skip_reasons: dict[str, int] = {}
    for entry in _load_registry(pipeline):
        doc_id = entry.get("document_id")
        ext = str(entry.get("extension", "")).lower()
        if not isinstance(doc_id, str) or not doc_id or doc_id in seen:
            skipped += 1
            skip_reasons["invalid_or_duplicate"] = skip_reasons.get("invalid_or_duplicate", 0) + 1
            _LOGGER.warning("Skipping invalid or duplicate persisted document entry")
            continue
        if ext not in SUPPORTED_EXTS:
            skipped += 1
            skip_reasons["unsupported_extension"] = skip_reasons.get("unsupported_extension", 0) + 1
            _LOGGER.warning("Skipping persisted document with unsupported extension")
            continue
        relative = entry.get("content_file", f"documents/{doc_id}.txt")
        content_path = (root / relative).resolve()
        if content_path.parent != documents_dir.resolve():
            skipped += 1
            skip_reasons["unsafe_content_reference"] = skip_reasons.get("unsafe_content_reference", 0) + 1
            _LOGGER.warning("Skipping persisted document with unsafe content reference")
            continue
        try:
            text = content_path.read_text(encoding="utf-8")
        except Exception:
            skipped += 1
            skip_reasons["read_failure"] = skip_reasons.get("read_failure", 0) + 1
            _LOGGER.exception("Failed to restore persisted document %s", doc_id)
            continue
        if not text.strip() or len(text) > MAX_EXTRACTED_TEXT_LEN:
            skipped += 1
            skip_reasons["empty_or_oversized"] = skip_reasons.get("empty_or_oversized", 0) + 1
            _LOGGER.warning("Skipping empty or oversized persisted document %s", doc_id)
            continue
        doc = UploadedDocument(
            name=str(entry.get("document_name") or "unnamed_document"),
            path=content_path,
            ext=ext,
            text=text,
            doc_id=doc_id,
            upload_timestamp=str(entry.get("upload_timestamp") or ""),
            source_type=str(entry.get("source_type") or "runtime_upload"),
            revision=entry.get("revision"),
        )
        doc.chunks = chunk_text(
            text, doc.doc_id, doc_name=doc.name, extension=doc.ext,
            upload_timestamp=doc.upload_timestamp, revision=doc.revision,
        )
        doc.chunk_count = len(doc.chunks)
        restored.append(doc)
        seen.add(doc_id)
    if skipped > 0 or restored:
        _LOGGER.info(
            "Runtime document recovery complete: restored=%d, skipped=%d, reasons=%s",
            len(restored), skipped, skip_reasons,
        )
    return restored


def _size_limit(path: Path) -> int:
    """Return the size limit for a given file based on its extension."""
    return UPLOAD_POLICY.per_file_bytes.get(path.suffix.lower(), 0)


def _validate_batch_limits(paths: list[Path]) -> str | None:
    total = sum(path.stat().st_size for path in paths)
    if total > UPLOAD_POLICY.max_batch_bytes:
        return (
            f"Upload batch exceeds total size limit "
            f"({UPLOAD_POLICY.max_batch_bytes} bytes)."
        )
    return None


def _sanitize_display_name(raw_name: str) -> str:
    """Produce a safe display filename from an untrusted raw name.

    - Extracts the basename (strips directory components).
    - Strips control characters.
    - Replaces path-traversal components.
    - Truncates to a reasonable length.
    - Never returns an empty string.
    """
    # Treat both separators as path delimiters regardless of host OS.
    name = re.split(r"[\\/]", str(raw_name))[-1]

    # Strip control characters (categories Cc, Cf)
    name = "".join(
        ch for ch in name
        if unicodedata.category(ch) not in ("Cc", "Cf")
    )

    # Remove any remaining path-traversal components
    name = name.replace("..", "").replace("~", "")

    # Collapse whitespace
    name = re.sub(r"\s+", " ", name).strip()

    # Truncate
    if len(name) > _MAX_DISPLAY_NAME_LEN:
        name = name[:_MAX_DISPLAY_NAME_LEN]

    # Fallback
    if not name:
        name = "unnamed_document"

    return name


@dataclass
class UploadedDocument:
    """Record of a single uploaded file with provenance metadata."""

    name: str
    path: Path
    ext: str
    text: str
    doc_id: str = field(default_factory=lambda: str(uuid.uuid4()))
    chunks: list[RuntimeChunk] = field(default_factory=list)
    chunk_count: int = 0
    upload_timestamp: str = field(default_factory=lambda: datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%S.%f")[:-3] + "Z")
    source_type: str = "runtime_upload"
    revision: str | None = None

    @property
    def safe_display_name(self) -> str:
        """Sanitized filename safe for public/API/UI display."""
        return _sanitize_display_name(self.name)

    def to_dict(self) -> dict:
        """Public-facing document summary.  No absolute paths."""
        return {
            "document_id": self.doc_id,
            "document_name": self.safe_display_name,
            "extension": self.ext,
            "chunk_count": self.chunk_count,
            "upload_timestamp": self.upload_timestamp,
            "source_type": self.source_type,
            "revision": self.revision,
        }


def _read_text(path: Path) -> str:
    """Plain TXT reader."""
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_pdf(path: Path) -> str:
    """PDF reader using PyPDF2 (lazy import)."""
    from PyPDF2 import PdfReader  # type: ignore

    reader = PdfReader(str(path))
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            parts.append("")
    return "\n\n".join(parts)


def _read_docx(path: Path) -> str:
    """DOCX reader using python-docx (lazy import)."""
    from docx import Document as DocxDocument  # type: ignore

    doc = DocxDocument(str(path))
    return "\n\n".join(p.text for p in doc.paragraphs)


def parse_file(path: Path) -> str:
    """Dispatch a single file to the right parser."""
    ext = path.suffix.lower()
    if ext == ".txt":
        return _read_text(path)
    if ext == ".pdf":
        return _read_pdf(path)
    if ext == ".docx":
        return _read_docx(path)
    raise ValueError(f"Unsupported file type: {ext}")


# ---------------------------------------------------------------------------
# Chunking
# ---------------------------------------------------------------------------

_WORD_RE = re.compile(r"\s+")
# Approximate chunk size in words; matches the mean chunk size produced by
# retriever_v2.load_chunks (~120 words/sentence, ~5 sentences per chunk).
_CHUNK_WORDS = 500
_OVERLAP_WORDS = 50


def _normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _WORD_RE.sub(" ", text)
    return text.strip()


def chunk_text(
    text: str,
    doc_id: str,
    *,
    doc_name: str = "",
    extension: str = ".txt",
    upload_timestamp: str = "",
    revision: str | None = None,
    chunk_words: int = _CHUNK_WORDS,
    overlap: int = _OVERLAP_WORDS,
) -> list[RuntimeChunk]:
    """Split text into overlapping word-window chunks with full provenance.

    Each returned ``RuntimeChunk`` carries a ``metadata`` dict containing:

    - document_id
    - document_name
    - chunk_index  (0-based, deterministic for identical input)
    - source_type  ("runtime_upload")
    - extension
    - upload_timestamp
    - page_number  (always None for this checkpoint)
    - revision     (always None unless explicitly provided)

    Args:
        text: The raw text to chunk.
        doc_id: Unique identifier for the source document.
        doc_name: Human-readable document name (sanitized for display).
        extension: File extension of the original document.
        upload_timestamp: ISO-8601 UTC timestamp of the upload.
        revision: Optional revision/version string.
        chunk_words: Approximate number of words per chunk.
        overlap: Number of overlapping words between consecutive chunks.
    """
    text = _normalize_text(text)
    if not text:
        return []

    safe_name = _sanitize_display_name(doc_name) if doc_name else "unnamed"

    def _make_chunk(piece: str, idx: int) -> RuntimeChunk:
        return RuntimeChunk(piece, metadata={
            "document_id": doc_id,
            "document_name": safe_name,
            "chunk_index": idx,
            "source_type": "runtime_upload",
            "extension": extension,
            "upload_timestamp": upload_timestamp,
            "page_number": None,
            "revision": revision,
        })

    words = text.split(" ")
    if len(words) <= chunk_words:
        return [_make_chunk(text, 0)]

    step = max(1, chunk_words - overlap)
    chunks: list[RuntimeChunk] = []
    idx = 0
    for start in range(0, len(words), step):
        piece_words = words[start : start + chunk_words]
        if not piece_words:
            break
        piece = " ".join(piece_words)
        chunks.append(_make_chunk(piece, idx))
        idx += 1
        if start + chunk_words >= len(words):
            break
    return chunks


# ---------------------------------------------------------------------------
# Pipeline integration
# ---------------------------------------------------------------------------

def _is_supported(path: Path) -> bool:
    return path.suffix.lower() in SUPPORTED_EXTS


def _attach_documents_unlocked(
    pipeline: dict,
    uploaded: list[UploadedDocument],
    *,
    persist: bool = True,
) -> int:
    """Merge uploaded chunks into the pipeline's chunks list and rebuild the index.

    Returns the total number of new chunks added.
    """
    new_chunks: list[RuntimeChunk] = []
    batch_bytes = 0
    existing_ids = {
        str(doc.get("document_id"))
        for doc in pipeline.get("uploaded_docs", [])
        if isinstance(doc, dict) and doc.get("document_id")
    }
    incoming_ids = [doc.doc_id for doc in uploaded]
    if len(incoming_ids) != len(set(incoming_ids)):
        raise ValueError("Duplicate document IDs are not allowed.")
    if existing_ids.intersection(incoming_ids):
        raise ValueError("Document ID already exists.")
    for doc in uploaded:
        if doc.ext.lower() not in UPLOAD_POLICY.allowed_extensions:
            raise ValueError(f"Unsupported file type: {doc.ext}")
        if len(doc.text) > UPLOAD_POLICY.max_extracted_text_chars:
            raise ValueError("Extracted text exceeds maximum allowed length.")
        if doc.path.exists():
            size_limit = _size_limit(doc.path)
            size = doc.path.stat().st_size
            if size_limit and size > size_limit:
                raise ValueError(f"File exceeds size limit for {doc.ext}.")
            batch_bytes += size
        if not doc.chunks:
            doc.chunks = chunk_text(
                doc.text,
                doc.doc_id,
                doc_name=doc.name,
                extension=doc.ext,
                upload_timestamp=doc.upload_timestamp,
                revision=doc.revision,
            )
        new_chunks.extend(doc.chunks)
        doc.chunk_count = len(doc.chunks)
        if doc.chunk_count > UPLOAD_POLICY.max_chunks_per_document:
            raise ValueError("Document exceeds maximum chunk limit.")
    if batch_bytes > UPLOAD_POLICY.max_batch_bytes:
        raise ValueError("Upload batch exceeds total size limit.")
    if len(new_chunks) > UPLOAD_POLICY.max_total_chunks_per_batch:
        raise ValueError("Upload batch exceeds maximum total chunk limit.")

    old_chunks = copy.deepcopy(pipeline.get("chunks", []))
    old_index = copy.deepcopy(pipeline.get("retrieval_index"))
    old_frequency = copy.deepcopy(pipeline.get("document_frequency"))
    old_uploaded_docs = list(pipeline.get("uploaded_docs", []))

    old_registry = None
    old_contents = {}
    if persist and pipeline.get("runtime_persistence", False):
        old_registry = _load_registry(pipeline)
        for doc in uploaded:
            content_path = _persistence_dir(pipeline) / "documents" / f"{doc.doc_id}.txt"
            if content_path.exists():
                try:
                    old_contents[doc.doc_id] = content_path.read_text(encoding="utf-8")
                except OSError:
                    old_contents[doc.doc_id] = None
            else:
                old_contents[doc.doc_id] = None

    try:
        if persist and pipeline.get("runtime_persistence", False):
            for doc in uploaded:
                _persist_document(pipeline, doc)
        if not new_chunks:
            return 0

        # retriever_v2 expects a list of RuntimeChunk objects.
        old_chunk_count = len(pipeline["chunks"])
        pipeline["chunks"].extend(new_chunks)
        index = pipeline.get("retrieval_index")
        frequency = pipeline.get("document_frequency")
        if (
            index is not None
            and frequency is not None
            and len(index) == old_chunk_count
        ):
            extend_index_v2(index, frequency, new_chunks, old_chunk_count)
        else:
            pipeline["retrieval_index"], pipeline["document_frequency"] = build_index_v2(
                pipeline["chunks"]
            )

        # Track uploads in the pipeline so the UI can list them.
        pipeline.setdefault("uploaded_docs", []).extend(
            [d.to_dict() for d in uploaded]
        )

        return len(new_chunks)
    except Exception:
        pipeline["chunks"] = old_chunks
        if old_index is not None:
            pipeline["retrieval_index"] = old_index
        if old_frequency is not None:
            pipeline["document_frequency"] = old_frequency
        pipeline["uploaded_docs"] = old_uploaded_docs

        if persist and pipeline.get("runtime_persistence", False) and old_registry is not None:
            try:
                _persist_registry(pipeline, old_registry)
            except Exception as registry_exc:
                _LOGGER.warning("Rollback: failed to restore registry: %s", type(registry_exc).__name__)
            for doc in uploaded:
                content_path = _persistence_dir(pipeline) / "documents" / f"{doc.doc_id}.txt"
                old_content = old_contents.get(doc.doc_id)
                try:
                    if old_content is not None:
                        _atomic_write(content_path, old_content)
                    else:
                        if content_path.exists():
                            content_path.unlink()
                except OSError as content_exc:
                    _LOGGER.warning("Rollback: failed to restore content for %s: %s", doc.doc_id, type(content_exc).__name__)

        raise


def attach_documents(
    pipeline: dict,
    uploaded: list[UploadedDocument],
    *,
    persist: bool = True,
) -> int:
    """Attach documents while serializing pipeline and registry mutations."""
    with _LIFECYCLE_LOCK:
        return _attach_documents(pipeline, uploaded, persist=persist)


def _attach_documents(
    pipeline: dict,
    uploaded: list[UploadedDocument],
    *,
    persist: bool = True,
) -> int:
    """Attach uploads after the public lifecycle lock is acquired."""
    return _attach_documents_unlocked(pipeline, uploaded, persist=persist)


def _remove_uploaded_document_unlocked(pipeline: dict, document_id: str) -> int:
    """Remove all RuntimeChunk objects belonging to the given document_id.

    - Removes only chunks whose ``metadata["document_id"]`` exactly matches.
    - Leaves static corpus chunks untouched.
    - Leaves other uploaded documents untouched.
    - Rebuilds the lexical index and updates ``pipeline["uploaded_docs"]``.
    - Returns the number of chunks removed.
    - If ``document_id`` is not found, returns 0 without modifying pipeline.
    """
    old_chunks = copy.deepcopy(pipeline.get("chunks", []))
    old_index = copy.deepcopy(pipeline.get("retrieval_index"))
    old_frequency = copy.deepcopy(pipeline.get("document_frequency"))
    old_uploaded_docs = list(pipeline.get("uploaded_docs", []))

    old_registry = None
    old_content = None
    content_path = _persistence_dir(pipeline) / "documents" / f"{document_id}.txt"
    if pipeline.get("runtime_persistence", False):
        old_registry = _load_registry(pipeline)
        if content_path.exists():
            try:
                old_content = content_path.read_text(encoding="utf-8")
            except OSError:
                old_content = None
        else:
            old_content = None

    chunks = pipeline.get("chunks", [])

    # Identify chunks to remove
    to_keep = []
    removed_count = 0
    for chunk in chunks:
        meta = getattr(chunk, "metadata", None)
        if (
            isinstance(meta, dict)
            and meta.get("document_id") == document_id
        ):
            removed_count += 1
        else:
            to_keep.append(chunk)

    try:
        if removed_count == 0:
            if pipeline.get("runtime_persistence", False):
                entries = _load_registry(pipeline)
                if any(d.get("document_id") == document_id for d in entries):
                    _persist_registry(
                        pipeline, [d for d in entries if d.get("document_id") != document_id]
                    )
                    content_path.unlink(missing_ok=True)
            return 0

        pipeline["chunks"] = to_keep
        pipeline["retrieval_index"], pipeline["document_frequency"] = build_index_v2(
            to_keep
        )

        # Remove from uploaded_docs tracking
        uploaded_docs = pipeline.get("uploaded_docs", [])
        pipeline["uploaded_docs"] = [
            d for d in uploaded_docs
            if d.get("document_id") != document_id
        ]
        if pipeline.get("runtime_persistence", False):
            entries = _load_registry(pipeline)
            _persist_registry(
                pipeline, [d for d in entries if d.get("document_id") != document_id]
            )
            content_path.unlink(missing_ok=True)

        return removed_count
    except Exception:
        pipeline["chunks"] = old_chunks
        if old_index is not None:
            pipeline["retrieval_index"] = old_index
        if old_frequency is not None:
            pipeline["document_frequency"] = old_frequency
        pipeline["uploaded_docs"] = old_uploaded_docs

        if pipeline.get("runtime_persistence", False) and old_registry is not None:
            try:
                _persist_registry(pipeline, old_registry)
            except Exception as registry_exc:
                _LOGGER.warning("Rollback: failed to restore registry: %s", type(registry_exc).__name__)
            try:
                if old_content is not None:
                    _atomic_write(content_path, old_content)
                else:
                    if content_path.exists():
                        content_path.unlink()
            except OSError as content_exc:
                _LOGGER.warning("Rollback: failed to restore content for %s: %s", document_id, type(content_exc).__name__)

        raise


def remove_uploaded_document(pipeline: dict, document_id: str) -> int:
    """Remove a document while serializing lifecycle and registry mutations."""
    with _LIFECYCLE_LOCK:
        return _remove_uploaded_document(pipeline, document_id)


def has_uploaded_document(pipeline: dict, document_id: str) -> bool:
    """Return whether an ID exists in runtime metadata or persisted registry."""
    with _LIFECYCLE_LOCK:
        if any(
            isinstance(doc, dict) and doc.get("document_id") == document_id
            for doc in pipeline.get("uploaded_docs", [])
        ):
            return True
        if pipeline.get("runtime_persistence", False):
            return any(
                entry.get("document_id") == document_id
                for entry in _load_registry(pipeline)
            )
        return False


def _remove_uploaded_document(
    pipeline: dict, document_id: str
) -> int:
    """Remove an upload after the public lifecycle lock is acquired."""
    return _remove_uploaded_document_unlocked(pipeline, document_id)


def process_uploads(
    pipeline: dict,
    file_paths: Iterable[str],
) -> tuple[list[UploadedDocument], list[str]]:
    """Parse a batch of uploaded files, return (parsed, errors).

    The pipeline is not modified here; call ``attach_documents`` after this
    step when the user explicitly confirms the upload.
    """
    parsed: list[UploadedDocument] = []
    errors: list[str] = []
    candidate_paths: list[Path] = []

    for raw in file_paths:
        path = Path(raw)
        if not path.exists():
            errors.append(f"Not found: {_sanitize_display_name(raw)}")
            continue
        if not _is_supported(path):
            errors.append(f"Unsupported file type: {_sanitize_display_name(path.name)}")
            continue
        candidate_paths.append(path)

    batch_error = _validate_batch_limits(candidate_paths)
    if batch_error:
        return [], [batch_error]

    for path in candidate_paths:
        # Size check
        size_limit = _size_limit(path)
        if size_limit and path.stat().st_size > size_limit:
            errors.append(
                f"File {_sanitize_display_name(path.name)} exceeds size limit "
                f"for {path.suffix.lower()} ({size_limit} bytes)"
            )
            continue

        # Parse file content — generic client error, details to logs only
        try:
            text = parse_file(path)
        except Exception:
            _LOGGER.exception("Failed to parse uploaded file: %s", path.name)
            errors.append(f"Unable to parse uploaded file.")
            continue

        # Empty or overly large extracted text check
        if not text.strip():
            errors.append(f"File {_sanitize_display_name(path.name)} contains no extractable text.")
            continue
        if len(text) > UPLOAD_POLICY.max_extracted_text_chars:
            errors.append(
                f"Extracted text from {_sanitize_display_name(path.name)} "
                f"exceeds maximum allowed length."
            )
            continue

        doc = UploadedDocument(
            name=path.name,
            path=path,
            ext=path.suffix.lower(),
            text=text,
        )
        doc.chunks = chunk_text(
            doc.text, doc.doc_id, doc_name=doc.name, extension=doc.ext,
            upload_timestamp=doc.upload_timestamp, revision=doc.revision,
        )
        doc.chunk_count = len(doc.chunks)
        if doc.chunk_count > UPLOAD_POLICY.max_chunks_per_document:
            errors.append(
                f"File {_sanitize_display_name(path.name)} exceeds maximum chunk limit."
            )
            continue
        parsed.append(doc)

    if sum(doc.chunk_count for doc in parsed) > UPLOAD_POLICY.max_total_chunks_per_batch:
        return [], [
            "Upload batch exceeds maximum total chunk limit "
            f"({UPLOAD_POLICY.max_total_chunks_per_batch} chunks)."
        ]
    return parsed, errors
